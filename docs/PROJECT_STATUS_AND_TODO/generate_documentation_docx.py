"""One-off script to build PROJECT_STATUS_AND_ARCHITECTURE.docx (requires python-docx)."""

from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

OUT = Path(__file__).resolve().parent / "PROJECT_STATUS_AND_ARCHITECTURE.docx"


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    run.font.size = Pt(11)
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main():
    doc = Document()
    t = doc.add_heading("TroyBanks Audit (Virginia) — Project Status & File Guide", 0)
    t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    add_para(
        doc,
        "This document summarizes work completed on the electricity tariff billing comparison "
        "platform, what is left to do before production use, and what each important file does. "
        "The project folder may be named ‘untitled folder 2’ or similar on your machine.",
    )

    # --- Completed ---
    add_heading(doc, "1. What we completed", level=1)
    add_bullets(
        doc,
        [
            "Fixed two billing-engine modules so they do not run PDF/OCR loops on import (required for a safe FastAPI startup).",
            "Added a FastAPI backend with REST endpoints for bills, calculations, anomalies, tariff/rider uploads, version history, and Excel export.",
            "Streamlit frontend (frontend/streamlit3.py) talks to the backend over HTTP instead of importing src business logic directly (except paths for upload directories).",
            "Added Dockerfiles for backend and frontend and a root docker-compose.yml that runs the app plus the existing Airflow stack.",
            "Removed the legacy monolithic src/Web_UI/streamlit3.py; use frontend/streamlit3.py (Docker) or src/Web_UI/streamlit.py (older local-only UI on va_step1_base.xlsx).",
        ],
    )

    # --- TODO ---
    add_heading(doc, "2. What you still need to do", level=1)
    add_para(doc, "Before production launch:", bold=False)
    add_bullets(
        doc,
        [
            "Ensure .env exists at the project root with DATABASE_URL (and Airflow variables if you use compose Airflow).",
            "First-time database: run create_all — e.g. docker compose exec backend python -c \"from src.Utils.createdb import *\"",
            "Build and start: docker compose up -d --build",
            "Open frontend (e.g. http://localhost:8501), backend health (http://localhost:8000/health), Airflow (http://localhost:8080) if enabled.",
            "On AWS EC2: open security group ports 22, 8501, 8000, 8080; install Docker; clone repo; configure .env; run docker compose.",
            "Optional: add reverse proxy (Nginx) on port 80, TLS certificates, and harden CORS (backend currently allows all origins).",
            "Optional: regenerate this Word file anytime with pip install python-docx and generate_documentation_docx.py.",
        ],
    )

    # --- Architecture ---
    add_heading(doc, "3. Target architecture (short)", level=1)
    add_para(
        doc,
        "Browser → Streamlit (frontend container) → HTTP → FastAPI (backend container) → PostgreSQL (e.g. AWS RDS). "
        "Business logic stays under src/. Heavy OCR and billing run in the backend. Airflow remains optional orchestration under infra/airflow/.",
    )

    # --- File guide ---
    add_heading(doc, "4. File-by-file guide", level=1)

    sections = [
        (
            "Root",
            [
                ("docker-compose.yml", "Runs backend, frontend, Postgres/Redis/Airflow services. Mounts ./src and ./data into backend and frontend for live code/data without rebuild."),
                ("README.md", "Existing project readme (not modified in this pass)."),
                ("FILE_GUIDE.md", "Existing file guide in repo where present."),
            ],
        ),
        (
            "backend/",
            [
                ("main.py", "FastAPI app: CORS, includes routers, GET /health."),
                ("billing_modules.py", "Loads new-bills_v2 and new-bills-profile via importlib (hyphenated filenames)."),
                ("usage_pipeline.py", " Builds usage DataFrame from pivoted extract + profile (aligned with former Streamlit logic)."),
                ("db_usage.py", "Queries usage_bill, merges batches, tariff/rider version helpers for calculations."),
                ("calc_service.py", "Resolves tariff workbook and riders, runs schedule functions via app_new, JSON-safe DataFrame export."),
                ("anomaly_service.py", "Builds anomaly table using process_troybanks_audit_data; mirrors former Streamlit anomaly pipeline."),
                ("routes/bills.py", "POST /api/bills/upload (PDF → OCR → DB). GET /api/bills list. GET /api/bills/{account} usage."),
                ("routes/calculate.py", "GET schedules and sources; POST /api/calculate for one or many schedule IDs."),
                ("routes/anomalies.py", "GET /api/anomalies/{account} with optional filters and threshold query params."),
                ("routes/tariffs.py", "POST /api/tariffs/upload and /api/riders/upload; copies workbooks to on-disk schedule/rider paths."),
                ("routes/versions.py", "GET /api/versions/{tariff_rates|rider_rates}."),
                ("routes/export.py", "POST /api/export → in-memory xlsx via xlsxwriter."),
                ("requirements.txt", "Backend Python dependencies including OCR-related packages."),
                ("Dockerfile", "Python 3.11 slim + tesseract/poppler; installs deps; runs uvicorn backend.main:app."),
            ],
        ),
        (
            "frontend/",
            [
                ("streamlit3.py", "Production UI (Docker): requests to BACKEND_URL; SCHEDULE_FUNCS proxy calls calculate API."),
                ("requirements.txt", "Streamlit, requests, pandas, etc. (lighter than backend)."),
                ("Dockerfile", "Slim image; runs streamlit on port 8501."),
            ],
        ),
        (
            "src/ (selected — no broad refactors)",
            [
                ("Billing_Engine/new-bills_v2.py", "Guarded bottom PDF batch loop with if __name__ == \"__main__\" so import is safe."),
                ("Billing_Engine/new-bills-profile.py", "Same guard for CLI/OCR driver block."),
                ("Billing_Engine/app_new.py", "Schedule functions and SCHEDULE_FUNCS; used by backend calculations."),
                ("Web_UI/streamlit.py", "Older local-only dashboard (streamlit run src/Web_UI/streamlit.py); uses on-disk va_step1_base.xlsx + riders. Not the Docker split stack."),
                ("Utils/upload.py", "DB upload/versioning for usage, tariffs, riders."),
                ("Utils/database.py", "SQLAlchemy engine from DATABASE_URL."),
                ("Utils/paths.py", "data/ and interim paths (NEW_BILLS_DIR, RIDERS_OUT, SCHEDULES_XLSX, etc.)."),
                ("va_step2_anomalies_db.py", "Core anomaly detection; called from backend anomaly_service."),
            ],
        ),
        (
            "infra/",
            [
                ("airflow/", "Unchanged per scope: DAGs, Dockerfile, former docker-compose still present; root compose supersedes for full stack."),
            ],
        ),
        (
            "This docs folder",
            [
                ("PROJECT_STATUS_AND_ARCHITECTURE.docx", "This Word document (generated output)."),
                ("generate_documentation_docx.py", "Script to regenerate the .docx; run with python-docx installed."),
            ],
        ),
    ]

    for title, files in sections:
        add_heading(doc, title, level=2)
        for name, desc in files:
            p = doc.add_paragraph()
            r = p.add_run(f"{name}: ")
            r.bold = True
            p.add_run(desc)
        doc.add_paragraph()

    add_heading(doc, "5. Regenerating this Word file", level=1)
    add_para(
        doc,
        "From the project root, create a venv, pip install python-docx, then run:",
    )
    add_para(doc, '  python docs/PROJECT_STATUS_AND_TODO/generate_documentation_docx.py')
    add_para(
        doc,
        "Output path: docs/PROJECT_STATUS_AND_TODO/PROJECT_STATUS_AND_ARCHITECTURE.docx",
    )

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
